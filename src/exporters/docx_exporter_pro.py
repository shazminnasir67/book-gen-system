"""
src/exporters/docx_exporter_pro.py
===================================
Professional DOCX export for the Automated Book Generation System.

Bugs fixed vs previous version:
  1. TOC showed outline sub-bullets (e.g. "• The shifting landscape…") because
     chapter.title stored raw outline text. Fixed with _resolve_chapter_title()
     which strips bullets, hashes, "Chapter N:" prefixes, and leading punctuation.
  2. Chapter heading was duplicated — exporter wrote "Chapter N: Title" then
     content started with the same heading. Fixed: _remove_duplicate_chapter_title
     now strips BOTH "Chapter N:" lines AND bare title matches.
  3. DOCX TOC field had placeholder text inside the field element which rendered
     as broken XML in some Word versions. Fixed: placeholder is now a separate
     paragraph outside the field, with a clear instruction.
  4. Copyright page had leading blank line. Fixed.
  5. Bold-only lines (e.g. "**Section Title**") were rendered as headings but
     still had the ** chars in the output. Fixed in _clean_markdown.
"""

import logging
import re
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from src.config import Config
from src.database.supabase_client import Book, Chapter

logger = logging.getLogger(__name__)


class DocxExporterPro:
    """Generates a professional DOCX file from book and chapter data."""

    def __init__(self, config: Config) -> None:
        self._config = config

    def export(self, book: Book, chapters: list[Chapter], output_path: str) -> None:
        try:
            doc = Document()
            self._setup_professional_styles(doc)

            self._add_cover_page(doc, book.title)
            self._add_page_break(doc)

            self._add_copyright_page(doc, book)
            self._add_page_break(doc)

            self._add_toc_page(doc)
            self._add_page_break(doc)

            for i, chapter in enumerate(chapters, 1):
                self._add_chapter_content(doc, chapter, i, book.title)
                self._add_page_break(doc)

            self._add_about_author(doc)
            self._setup_headers_footers(doc, book.title)

            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            logger.info("DOCX saved: %s", output_path)
        except Exception as exc:
            logger.error("DOCX export failed for '%s': %s", book.title, exc)
            raise RuntimeError(f"DOCX export failed: {exc}") from exc

    # ── Title resolution (THE key fix) ────────────────────────────────────

    def _resolve_chapter_title(self, raw_title: str, content: str, chapter_num: int, book_title: str) -> str:
        """
        Return a clean chapter title regardless of what is stored in the DB.

        Handles these bad inputs:
          • "• The shifting landscape of modern medicine: From reactive to proactive care."
          • "Chapter 1: Building the Foundation: Strategy and Governance"
          • "Digital Transformation in Healthcare: A Practical Guide"  (= book title)
          • "## Introduction: The Imperative for Digital Evolution"
        """
        # Step 1: clean markdown / bullet chars
        title = self._clean_markdown(raw_title).strip()

        # Step 2: strip leading bullet characters and punctuation
        title = re.sub(r'^[•\-\*\#\s]+', '', title).strip()

        # Step 3: strip "Chapter N:" prefix (any chapter number)
        title = re.sub(r'^Chapter\s+\d+\s*:\s*', '', title, flags=re.IGNORECASE).strip()

        # Step 4: if what's left is empty, same as book title, or very short → mine from content
        clean_book = self._clean_markdown(book_title).strip().lower()
        if not title or title.lower() == clean_book or len(title) < 4:
            title = self._extract_title_from_content(content, chapter_num)

        return title or f"Chapter {chapter_num}"

    def _extract_title_from_content(self, content: str, chapter_num: int) -> str:
        """Pull the first real heading out of content that isn't a bare chapter number line."""
        for line in content.splitlines():
            line = line.strip()
            if not line:
                continue
            # Skip lines that are just "Chapter N" or "Chapter N:"
            if re.match(r'^Chapter\s+\d+\s*:?\s*$', line, re.IGNORECASE):
                continue
            # Prefer markdown heading lines
            m = re.match(r'^#{1,6}\s+(.+)$', line)
            if m:
                candidate = self._clean_markdown(m.group(1)).strip()
                candidate = re.sub(r'^Chapter\s+\d+\s*:\s*', '', candidate, flags=re.IGNORECASE).strip()
                if candidate and len(candidate) > 4:
                    return candidate
            # Or a bold-only line
            m2 = re.match(r'^\*\*(.+?)\*\*$', line)
            if m2:
                candidate = m2.group(1).strip()
                if candidate and len(candidate) > 4:
                    return candidate
        return ""

    # ── Styles ─────────────────────────────────────────────────────────────

    def _setup_professional_styles(self, doc: Document) -> None:
        for section in doc.sections:
            section.top_margin    = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin   = Inches(1.5)
            section.right_margin  = Inches(1.0)
            section.header_distance = Inches(0.5)
            section.footer_distance = Inches(0.5)

        style = doc.styles["Normal"]
        style.font.name = "Garamond"
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(6)

    # ── Cover ──────────────────────────────────────────────────────────────

    def _add_cover_page(self, doc: Document, title: str) -> None:
        for _ in range(10):
            doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title.upper())
        r.bold = True
        r.font.name = "Garamond"
        r.font.size = Pt(32)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        doc.add_paragraph()

        line_p = doc.add_paragraph()
        line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = line_p.add_run("═" * 30)
        lr.font.size = Pt(14)
        lr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_paragraph()

        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sub_p.add_run("Generated by Automated Book Generation System")
        sr.font.name = "Garamond"
        sr.font.size = Pt(14)
        sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        sr.italic = True

        for _ in range(8):
            doc.add_paragraph()

        year_p = doc.add_paragraph()
        year_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        yr = year_p.add_run(str(datetime.now().year))
        yr.font.name = "Garamond"
        yr.font.size = Pt(16)

    # ── Copyright ──────────────────────────────────────────────────────────

    def _add_copyright_page(self, doc: Document, book: Book) -> None:
        for _ in range(15):
            doc.add_paragraph()

        lines = [
            f"Copyright © {datetime.now().year}",
            "",
            "This book was automatically generated using AI technology.",
            "All rights reserved.",
            "",
            f"Generated: {datetime.now().strftime('%B %d, %Y')}",
            "System: Automated Book Generation System",
            "Model: Google Gemini",
            "Database: Supabase",
            "",
            "For more information about this technology:",
            "contact@example.com",
        ]
        for line in lines:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = "Garamond"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── TOC ────────────────────────────────────────────────────────────────

    def _add_toc_page(self, doc: Document) -> None:
        """Add TOC page with a proper Word field code (no text inside the field)."""
        toc_h = doc.add_paragraph("Table of Contents")
        toc_h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = toc_h.runs[0]
        r.font.name = "Garamond"
        r.font.size = Pt(24)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        r.bold = True

        doc.add_paragraph()

        # TOC field — no stray text inside the field
        para = doc.add_paragraph()
        run = para.add_run()

        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

        fldChar_sep = OxmlElement("w:fldChar")
        fldChar_sep.set(qn("w:fldCharType"), "separate")

        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")

        run._r.append(fldChar_begin)
        run._r.append(instrText)
        run._r.append(fldChar_sep)
        run._r.append(fldChar_end)

        # Instruction paragraph OUTSIDE the field — always visible
        hint = doc.add_paragraph()
        hint.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hint_run = hint.add_run(
            "To update the Table of Contents: press Ctrl+A then F9 (Windows) "
            "or Cmd+A then Fn+F9 (Mac), then choose 'Update entire table'."
        )
        hint_run.font.name = "Garamond"
        hint_run.font.size = Pt(10)
        hint_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        hint_run.italic = True

    # ── Chapter content ────────────────────────────────────────────────────

    def _add_chapter_content(self, doc: Document, chapter: Chapter, number: int, book_title: str) -> None:
        content = chapter.content or ""
        clean_title = self._resolve_chapter_title(
            chapter.title or "", content, number, book_title
        )

        # Heading 1 — appears in Word TOC
        heading = doc.add_heading(f"Chapter {number}: {clean_title}", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in heading.runs:
            run.font.name = "Garamond"
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            run.bold = True

        doc.add_paragraph()

        # Strip duplicate title lines from content before rendering
        content = self._remove_duplicate_title(content, number, clean_title)
        self._parse_and_add_content(doc, content)

    def _remove_duplicate_title(self, content: str, chapter_num: int, clean_title: str) -> str:
        """
        Remove lines from content that repeat the chapter heading.
        Handles: 'Chapter N: Title', '# Title', bare 'Title', '## Title'.
        """
        lines = content.splitlines()
        out = []
        title_lower = clean_title.lower().strip()

        for line in lines:
            stripped = line.strip()
            # Remove "Chapter N: ..." prefix variants
            if re.match(rf'^[#\s]*Chapter\s+{chapter_num}\s*:', stripped, re.IGNORECASE):
                continue
            # Remove bare title match (with or without leading #)
            bare = re.sub(r'^#{1,6}\s*', '', stripped).strip().lower()
            if bare == title_lower and len(bare) > 4:
                continue
            out.append(line)
        return "\n".join(out)

    def _parse_and_add_content(self, doc: Document, content: str) -> None:
        lines = content.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue

            # Markdown headings
            m = re.match(r'^(#{1,6})\s+(.+)$', line)
            if m:
                level = len(m.group(1))
                text = self._clean_markdown(m.group(2))
                heading_level = min(level + 1, 9)
                h = doc.add_heading(text, level=heading_level)
                h.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in h.runs:
                    run.font.name = "Garamond"
                    run.font.size = Pt(18 if level == 1 else 16 if level == 2 else 14)
                    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                i += 1
                continue

            # Standalone **bold line** → sub-heading
            m2 = re.match(r'^\*\*(.+?)\*\*$', line)
            if m2 and len(line) < 150:
                text = m2.group(1).strip()
                h = doc.add_heading(text, level=3)
                h.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in h.runs:
                    run.font.name = "Garamond"
                    run.font.size = Pt(15)
                    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                    run.bold = True
                i += 1
                continue

            # Bullet list
            b = re.match(r'^[-•]\s+(.+)$', line) or re.match(r'^\*\s+(.+)$', line)
            if b:
                while i < len(lines):
                    bl = lines[i].strip()
                    bc = re.match(r'^[-•]\s+(.+)$', bl) or re.match(r'^\*\s+(.+)$', bl)
                    if bc:
                        p = doc.add_paragraph(self._clean_markdown(bc.group(1)), style="List Bullet")
                        for run in p.runs:
                            run.font.name = "Garamond"
                            run.font.size = Pt(12)
                        p.paragraph_format.space_after = Pt(3)
                        i += 1
                    else:
                        break
                continue

            # Regular prose — collect paragraph lines
            para_lines = [line]
            i += 1
            while i < len(lines):
                nl = lines[i].strip()
                if not nl:
                    break
                if re.match(r'^#{1,6}\s+', nl):
                    break
                if re.match(r'^[-•]\s+', nl) or re.match(r'^\*\s+', nl):
                    break
                para_lines.append(nl)
                i += 1

            text = self._clean_markdown(" ".join(para_lines))
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = "Garamond"
                run.font.size = Pt(12)
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_after = Pt(6)

    # ── About Author ───────────────────────────────────────────────────────

    def _add_about_author(self, doc: Document) -> None:
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hr = h.add_run("About the Author")
        hr.font.name = "Garamond"
        hr.font.size = Pt(20)
        hr.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        hr.bold = True

        doc.add_paragraph()

        bio = (
            "Shazmin Nasir is a Data Scientist, AI Engineer, and UI/UX Designer "
            "passionate about automation from design to deployment. With expertise in "
            "artificial intelligence and user experience, Shazmin builds systems that "
            "bridge the gap between intelligent algorithms and intuitive interfaces, "
            "creating seamless solutions that delight users while leveraging cutting-edge technology."
        )
        bp = doc.add_paragraph(bio)
        bp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in bp.runs:
            run.font.name = "Garamond"
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

        doc.add_paragraph()

        cp = doc.add_paragraph()
        cr = cp.add_run("Contact: shazminnasir481@gmail.com")
        cr.font.name = "Garamond"
        cr.font.size = Pt(11)
        cr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        cr.italic = True

    # ── Headers / Footers ──────────────────────────────────────────────────

    def _setup_headers_footers(self, doc: Document, book_title: str) -> None:
        for section in doc.sections:
            header_para = section.header.paragraphs[0]
            header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            hr = header_para.add_run(book_title[:60])
            hr.font.name = "Garamond"
            hr.font.size = Pt(10)
            hr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            hr.italic = True

            footer_para = section.footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fr = footer_para.add_run()
            fld = OxmlElement("w:fldSimple")
            fld.set(qn("w:instr"), "PAGE")
            fr._r.append(fld)
            fr.font.name = "Garamond"
            fr.font.size = Pt(10)

    # ── Utilities ──────────────────────────────────────────────────────────

    def _clean_markdown(self, text: str) -> str:
        """Strip all markdown formatting, returning plain text."""
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)   # bold
        text = re.sub(r'\*(.*?)\*',   r'\1', text)      # italic
        text = re.sub(r'_(.*?)_',     r'\1', text)      # italic underscore
        text = re.sub(r'^#{1,6}\s+',  '',    text, flags=re.MULTILINE)  # headers
        text = re.sub(r'```.*?```',   '',    text, flags=re.DOTALL)     # code blocks
        text = re.sub(r'`(.*?)`',     r'\1', text)      # inline code
        text = re.sub(r'^[-*•]\s+',   '',    text, flags=re.MULTILINE)  # leading bullets
        return text.strip()

    def _add_page_break(self, doc: Document) -> None:
        doc.add_page_break()